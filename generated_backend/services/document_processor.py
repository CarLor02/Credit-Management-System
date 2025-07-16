"""
文档处理服务
处理已上传的文档，使用OCR工具进行解析
"""

import os
import sys
import uuid
import threading
import subprocess
from datetime import datetime
from pathlib import Path
from typing import Optional, Dict, Any

# 导入文档处理库
try:
    from docx import Document as DocxDocument
    import docx2txt
    import mammoth
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

from flask import current_app
from database import db
from db_models import Document, DocumentStatus

class DocumentProcessor:
    """文档处理器"""
    
    def __init__(self):
        self.ocr_path = self._get_ocr_path()
        self.processed_folder = 'processed'
        
    def _get_ocr_path(self) -> str:
        """获取OCR代码路径"""
        # 假设OCR文件夹在项目根目录下
        current_dir = os.path.dirname(os.path.abspath(__file__))
        project_root = os.path.dirname(os.path.dirname(current_dir))
        return os.path.join(project_root, 'OCR')
    
    def process_document_async(self, document_id: int):
        """异步处理文档"""
        def process_in_background():
            try:
                with current_app.app_context():
                    self.process_document(document_id)
            except Exception as e:
                current_app.logger.error(f"后台处理文档失败: {e}")
        
        thread = threading.Thread(target=process_in_background)
        thread.daemon = True
        thread.start()
    
    def process_document(self, document_id: int) -> bool:
        """处理单个文档"""
        try:
            document = Document.query.get(document_id)
            if not document:
                current_app.logger.error(f"文档不存在: {document_id}")
                return False
            
            # 更新状态为处理中，进度为10%
            document.status = DocumentStatus.PROCESSING
            document.processing_started_at = datetime.utcnow()
            document.progress = 10
            db.session.commit()
            
            current_app.logger.info(f"开始处理文档: {document.filename}")
            
            # 构建输入和输出路径
            input_file = document.file_path
            if not os.path.exists(input_file):
                current_app.logger.error(f"源文件不存在: {input_file}")
                self._mark_processing_failed(document, "源文件不存在")
                return False
            
            # 创建处理后文件的路径结构，进度为20%
            processed_file_path = self._create_processed_file_path(document)
            document.progress = 20
            db.session.commit()
            
            # 调用OCR处理，进度为30%
            document.progress = 30
            db.session.commit()
            
            success = self._call_ocr_processor(input_file, processed_file_path, document)
            
            if success:
                # 更新文档状态为已完成，进度为100%
                document.status = DocumentStatus.COMPLETED
                document.processed_file_path = processed_file_path
                document.processed_at = datetime.utcnow()
                document.progress = 100
                db.session.commit()
                
                current_app.logger.info(f"文档处理完成: {document.filename}")
                
                # 检查是否需要创建知识库
                self._check_and_create_knowledge_base(document)
                
                return True
            else:
                self._mark_processing_failed(document, "OCR处理失败")
                return False
                
        except Exception as e:
            current_app.logger.error(f"处理文档失败: {e}")
            if 'document' in locals():
                self._mark_processing_failed(document, str(e))
            return False
    
    def _create_processed_file_path(self, document: Document) -> str:
        """创建处理后文件的路径"""
        # 获取项目的folder_uuid
        project = document.project
        if not project or not project.folder_uuid:
            raise ValueError("项目不存在或缺少folder_uuid")
        
        # 创建processed目录结构：processed/{project_folder_uuid}/{filename}.md
        # processed和uploads应该是同级关系
        project_upload_dir = os.path.dirname(document.file_path)  # 获取uploads/{project_folder_uuid}目录
        uploads_dir = os.path.dirname(project_upload_dir)  # 获取uploads目录
        backend_dir = os.path.dirname(uploads_dir)  # 获取generated_backend目录
        
        processed_dir = os.path.join(
            backend_dir,
            self.processed_folder,
            project.folder_uuid
        )
        
        # 确保目录存在
        os.makedirs(processed_dir, exist_ok=True)
        
        # 获取上传时生成的文件名（带UUID前缀），去掉扩展名后加上.md
        uploaded_filename = os.path.basename(document.file_path)  # 获取带UUID的文件名
        base_name = os.path.splitext(uploaded_filename)[0]  # 去掉扩展名
        processed_filename = f"{base_name}.md"  # 加上.md扩展名
        
        return os.path.join(processed_dir, processed_filename)
    
    def _call_ocr_processor(self, input_file: str, output_file: str, document: Document) -> bool:
        """调用OCR处理器或直接处理特定文件类型"""
        try:
            # 确保输出目录存在
            output_dir = os.path.dirname(output_file)
            os.makedirs(output_dir, exist_ok=True)
            
            # 检测文件类型，进度为40%
            file_extension = Path(input_file).suffix.lower()
            file_type = self._get_file_type(file_extension)
            
            if not file_type:
                current_app.logger.error(f"不支持的文件类型: {file_extension}")
                return False
            
            document.progress = 40
            db.session.commit()
            
            # 对于markdown和word文件，直接调用专门的处理方法
            if file_type == 'markdown':
                return self._process_markdown_direct(input_file, output_file, document)
            elif file_type == 'word':
                return self._process_word_direct(input_file, output_file, document)
            else:
                # 其他文件类型使用OCR处理
                return self._process_with_ocr(input_file, output_file, document, file_type)
                
        except Exception as e:
            current_app.logger.error(f"调用处理器失败: {e}")
            return False
    
    def _process_markdown_direct(self, input_file: str, output_file: str, document: Document) -> bool:
        """直接处理markdown文件"""
        try:
            # 进度为50%
            document.progress = 50
            db.session.commit()
            
            # 读取原始md文件内容
            with open(input_file, 'r', encoding='utf-8') as f:
                content = f.read()
            
            # 进度为80%
            document.progress = 80
            db.session.commit()
            
            # 写入到processed目录
            with open(output_file, 'w', encoding='utf-8') as f:
                f.write(content)
            
            current_app.logger.info(f"markdown文件直接处理完成: {output_file}")
            return True
            
        except Exception as e:
            current_app.logger.error(f"直接处理markdown文件失败: {e}")
            return False
    
    def _process_word_direct(self, input_file: str, output_file: str, document: Document) -> bool:
        """直接处理Word文件"""
        try:
            # 检查是否有docx处理库
            if not DOCX_AVAILABLE:
                current_app.logger.error("缺少docx处理库，请安装 python-docx、docx2txt 和 mammoth")
                return False
            
            # 进度为50%
            document.progress = 50
            db.session.commit()
            
            # 提取Word文件内容
            markdown_content = self._extract_word_content(input_file, document.filename)
            
            if not markdown_content:
                current_app.logger.error("无法提取Word文件内容")
                return False
            
            # 进度为80%
            document.progress = 80
            db.session.commit()
            
            # 写入到processed目录
            with open(output_file, 'w', encoding='utf-8') as f:
                f.write(markdown_content)
            
            current_app.logger.info(f"Word文件直接处理完成: {output_file}")
            return True
            
        except Exception as e:
            current_app.logger.error(f"直接处理Word文件失败: {e}")
            return False
    
    def _process_with_ocr(self, input_file: str, output_file: str, document: Document, file_type: str) -> bool:
        """使用OCR处理文件"""
        try:
            # 获取输出目录
            output_dir = os.path.dirname(output_file)
            
            # 构建OCR处理命令
            ocr_script = os.path.join(self.ocr_path, 'process_docs.py')
            input_dir = os.path.dirname(input_file)
            
            # 使用绝对路径
            input_dir = os.path.abspath(input_dir)
            temp_output_dir = os.path.abspath(os.path.join(output_dir, f"temp_{uuid.uuid4().hex[:8]}"))
            
            # 使用单独的输出目录避免冲突
            
            cmd = [
                sys.executable, ocr_script,
                '--input', input_dir,
                '--output', temp_output_dir,
                '--type', file_type,
                '--no-recursive',
                '--no-preserve-structure'
            ]
            
            # 添加API密钥（如果需要）
            api_key = os.environ.get('YUNWU_API_KEY')
            if api_key:
                cmd.extend(['--api-key', api_key])
            
            current_app.logger.info(f"执行OCR命令: {' '.join(cmd)}")
            
            # 开始OCR处理，进度为50%
            document.progress = 50
            db.session.commit()
            
            # 执行OCR处理
            result = subprocess.run(
                cmd,
                capture_output=True,
                text=True,
                timeout=300,  # 5分钟超时
                cwd=self.ocr_path
            )
            
            # OCR处理完成，进度为80%
            document.progress = 80
            db.session.commit()
            
            current_app.logger.info(f"OCR命令退出码: {result.returncode}")
            if result.stdout:
                current_app.logger.info(f"OCR标准输出: {result.stdout}")
            if result.stderr:
                current_app.logger.error(f"OCR错误输出: {result.stderr}")
            
            if result.returncode == 0:
                # OCR成功，查找生成的文件并移动到目标位置
                processed_files = list(Path(temp_output_dir).rglob('*.md'))
                
                current_app.logger.info(f"在{temp_output_dir}中找到{len(processed_files)}个处理后的文件")
                
                if processed_files:
                    # 查找与当前文档匹配的文件
                    target_file = None
                    base_filename = os.path.splitext(document.filename)[0]
                    
                    # 首先尝试找到与文档名称匹配的文件
                    for file in processed_files:
                        if base_filename in file.stem:
                            target_file = file
                            break
                    
                    # 如果没有找到匹配的文件，取第一个
                    if target_file is None:
                        target_file = processed_files[0]
                    
                    current_app.logger.info(f"移动文件: {target_file} -> {output_file}")
                    
                    # 移动文件到最终位置，进度为90%
                    document.progress = 90
                    db.session.commit()
                    
                    # 移动文件到最终位置
                    import shutil
                    shutil.move(str(target_file), output_file)
                    
                    # 清理临时目录
                    shutil.rmtree(temp_output_dir, ignore_errors=True)
                    
                    current_app.logger.info(f"OCR处理成功，输出文件: {output_file}")
                    return True
                else:
                    current_app.logger.error("OCR处理未生成任何文件")
                    return False
            else:
                current_app.logger.error(f"OCR处理失败: {result.stderr}")
                return False
                
        except subprocess.TimeoutExpired:
            current_app.logger.error("OCR处理超时")
            return False
        except Exception as e:
            current_app.logger.error(f"调用OCR处理器失败: {e}")
            return False
    
    def _get_file_type(self, extension: str) -> Optional[str]:
        """根据文件扩展名获取处理类型"""
        extension = extension.lower()
        
        if extension == '.pdf':
            return 'pdf'
        elif extension in ['.xlsx', '.xls', '.csv']:
            return 'excel'
        elif extension in ['.html', '.htm']:
            return 'html'
        elif extension in ['.jpg', '.jpeg', '.png', '.gif', '.webp', '.bmp', '.tiff', '.tif']:
            return 'image'
        elif extension in ['.md', '.markdown']:
            return 'markdown'
        elif extension in ['.doc', '.docx']:
            return 'word'
        else:
            return None
    
    def _mark_processing_failed(self, document: Document, error_message: str):
        """标记文档处理失败"""
        try:
            document.status = DocumentStatus.FAILED
            document.progress = 0  # 失败时进度设为0
            db.session.commit()
            current_app.logger.error(f"文档处理失败: {document.filename}, 错误: {error_message}")
        except Exception as e:
            current_app.logger.error(f"更新文档状态失败: {e}")
            db.session.rollback()
    
    def get_processing_progress(self, document_id: int) -> Dict[str, Any]:
        """获取文档处理进度"""
        try:
            document = Document.query.get(document_id)
            if not document:
                return {'error': '文档不存在'}
            
            progress_info = {
                'status': document.status.value,
                'filename': document.filename,
                'processing_started_at': document.processing_started_at.isoformat() if document.processing_started_at else None,
                'processed_at': document.processed_at.isoformat() if document.processed_at else None,
                'processed_file_path': document.processed_file_path,
                'progress_percent': document.progress  # 使用数据库中的实际进度
            }
            
            return progress_info
            
        except Exception as e:
            current_app.logger.error(f"获取处理进度失败: {e}")
            return {'error': '获取进度失败'}
    
    def delete_processed_document(self, document: Document) -> bool:
        """删除处理后的文档文件"""
        try:
            if document.processed_file_path and os.path.exists(document.processed_file_path):
                os.remove(document.processed_file_path)
                current_app.logger.info(f"删除处理后的文档: {document.processed_file_path}")
                return True
            return False
        except Exception as e:
            current_app.logger.error(f"删除处理后的文档失败: {e}")
            return False
    
    def delete_project_documents(self, project) -> bool:
        """删除项目相关的所有文档文件（原始文件和处理后文件）"""
        try:
            success_count = 0
            total_count = 0
            
            # 获取项目的所有文档
            documents = Document.query.filter_by(project_id=project.id).all()
            
            for document in documents:
                total_count += 1
                
                # 删除原始文件
                try:
                    if document.file_path and os.path.exists(document.file_path):
                        os.remove(document.file_path)
                        current_app.logger.info(f"删除原始文件: {document.file_path}")
                except Exception as e:
                    current_app.logger.warning(f"删除原始文件失败: {e}")
                
                # 删除处理后的文件
                try:
                    if document.processed_file_path and os.path.exists(document.processed_file_path):
                        os.remove(document.processed_file_path)
                        current_app.logger.info(f"删除处理后文件: {document.processed_file_path}")
                except Exception as e:
                    current_app.logger.warning(f"删除处理后文件失败: {e}")
                
                success_count += 1
            
            # 删除项目相关的文件夹
            try:
                # 删除uploads中的项目文件夹
                upload_folder = current_app.config.get('UPLOAD_FOLDER', 'uploads')
                project_upload_dir = os.path.join(upload_folder, project.folder_uuid)
                if os.path.exists(project_upload_dir):
                    import shutil
                    shutil.rmtree(project_upload_dir)
                    current_app.logger.info(f"删除项目上传目录: {project_upload_dir}")
                
                # 删除processed中的项目文件夹
                backend_dir = os.path.dirname(upload_folder)
                project_processed_dir = os.path.join(backend_dir, self.processed_folder, project.folder_uuid)
                if os.path.exists(project_processed_dir):
                    import shutil
                    shutil.rmtree(project_processed_dir)
                    current_app.logger.info(f"删除项目处理目录: {project_processed_dir}")
                    
            except Exception as e:
                current_app.logger.warning(f"删除项目文件夹失败: {e}")
            
            current_app.logger.info(f"项目文档删除完成: {success_count}/{total_count}")
            return success_count == total_count
            
        except Exception as e:
            current_app.logger.error(f"删除项目文档失败: {e}")
            return False
    
    def process_markdown_file(self, document_id: int) -> bool:
        """处理markdown文件（无需OCR，直接复制）"""
        try:
            document = Document.query.get(document_id)
            if not document:
                current_app.logger.error(f"文档不存在: {document_id}")
                return False
            
            # 更新状态为处理中，进度为10%
            document.status = DocumentStatus.PROCESSING
            document.processing_started_at = datetime.utcnow()
            document.progress = 10
            db.session.commit()
            
            current_app.logger.info(f"开始处理markdown文件: {document.filename}")
            
            # 构建输入和输出路径
            input_file = document.file_path
            if not os.path.exists(input_file):
                current_app.logger.error(f"源文件不存在: {input_file}")
                self._mark_processing_failed(document, "源文件不存在")
                return False
            
            # 创建处理后文件的路径，进度为30%
            processed_file_path = self._create_processed_file_path(document)
            document.progress = 30
            db.session.commit()
            
            # 直接复制md文件内容到processed目录，进度为70%
            document.progress = 70
            db.session.commit()
            
            # 读取原始md文件内容
            with open(input_file, 'r', encoding='utf-8') as f:
                content = f.read()
            
            # 写入到processed目录
            with open(processed_file_path, 'w', encoding='utf-8') as f:
                f.write(content)
            
            # 更新文档状态为已完成，进度为100%
            document.status = DocumentStatus.COMPLETED
            document.processed_file_path = processed_file_path
            document.processed_at = datetime.utcnow()
            document.progress = 100
            db.session.commit()
            
            current_app.logger.info(f"markdown文件处理完成: {document.filename}")
            
            # 检查是否需要创建知识库
            self._check_and_create_knowledge_base(document)
            
            return True
                
        except Exception as e:
            current_app.logger.error(f"处理markdown文件失败: {e}")
            if 'document' in locals():
                self._mark_processing_failed(document, str(e))
            return False
    
    def process_word_file(self, document_id: int) -> bool:
        """处理Word文件（doc/docx）"""
        try:
            document = Document.query.get(document_id)
            if not document:
                current_app.logger.error(f"文档不存在: {document_id}")
                return False
            
            # 检查是否有docx处理库
            if not DOCX_AVAILABLE:
                current_app.logger.error("缺少docx处理库，请安装 python-docx、docx2txt 和 mammoth")
                self._mark_processing_failed(document, "缺少docx处理库")
                return False
            
            # 更新状态为处理中，进度为10%
            document.status = DocumentStatus.PROCESSING
            document.processing_started_at = datetime.utcnow()
            document.progress = 10
            db.session.commit()
            
            current_app.logger.info(f"开始处理Word文件: {document.filename}")
            
            # 构建输入和输出路径
            input_file = document.file_path
            if not os.path.exists(input_file):
                current_app.logger.error(f"源文件不存在: {input_file}")
                self._mark_processing_failed(document, "源文件不存在")
                return False
            
            # 创建处理后文件的路径，进度为30%
            processed_file_path = self._create_processed_file_path(document)
            document.progress = 30
            db.session.commit()
            
            # 提取Word文件内容，进度为50%
            document.progress = 50
            db.session.commit()
            
            markdown_content = self._extract_word_content(input_file, document.filename)
            
            if not markdown_content:
                self._mark_processing_failed(document, "无法提取Word文件内容")
                return False
            
            # 进度为80%
            document.progress = 80
            db.session.commit()
            
            # 写入到processed目录
            with open(processed_file_path, 'w', encoding='utf-8') as f:
                f.write(markdown_content)
            
            # 更新文档状态为已完成，进度为100%
            document.status = DocumentStatus.COMPLETED
            document.processed_file_path = processed_file_path
            document.processed_at = datetime.utcnow()
            document.progress = 100
            db.session.commit()
            
            current_app.logger.info(f"Word文件处理完成: {document.filename}")
            
            # 检查是否需要创建知识库
            self._check_and_create_knowledge_base(document)
            
            return True
                
        except Exception as e:
            current_app.logger.error(f"处理Word文件失败: {e}")
            if 'document' in locals():
                self._mark_processing_failed(document, str(e))
            return False
    
    def _extract_word_content(self, file_path: str, filename: str) -> str:
        """提取Word文件内容并转换为Markdown格式"""
        try:
            file_extension = Path(file_path).suffix.lower()
            
            if file_extension == '.docx':
                # 处理.docx文件，优先使用mammoth，fallback到python-docx
                return self._extract_docx_content(file_path, filename)
            elif file_extension == '.doc':
                # 处理.doc文件
                return self._extract_doc_content(file_path, filename)
            else:
                current_app.logger.error(f"不支持的文件类型: {file_extension}")
                return self._generate_doc_conversion_message(f"不支持的文件类型: {file_extension}")
                
        except Exception as e:
            current_app.logger.error(f"提取Word内容失败: {e}")
            return self._generate_doc_conversion_message("文档处理过程中发生错误", str(e))
    
    def _extract_docx_content(self, file_path: str, filename: str) -> str:
        """提取.docx文件内容，优先使用mammoth，fallback到python-docx"""
        try:
            # 方法1：使用mammoth（更好的markdown转换）
            try:
                with open(file_path, "rb") as docx_file:
                    result = mammoth.convert_to_markdown(docx_file)
                    if result.value and result.value.strip():
                        # 添加文档标题
                        document_title = Path(filename).stem
                        markdown_content = f"# {document_title}\n\n{result.value}"
                        
                        # 如果有警告信息，添加到文档末尾
                        if result.messages:
                            warnings = [msg.message for msg in result.messages]
                            markdown_content += f"\n\n---\n*转换警告: {'; '.join(warnings)}*"
                        
                        current_app.logger.info(f"使用mammoth成功转换docx文件: {filename}")
                        return markdown_content
                    else:
                        current_app.logger.warning(f"mammoth转换结果为空，使用fallback方法")
                        return self._extract_docx_with_python_docx(file_path, filename)
            except Exception as e:
                current_app.logger.warning(f"mammoth转换失败: {e}，使用fallback方法")
                return self._extract_docx_with_python_docx(file_path, filename)
                
        except Exception as e:
            current_app.logger.error(f"提取.docx内容失败: {e}")
            return self._generate_doc_conversion_message("DOCX文档处理失败", str(e))
    
    def _extract_docx_with_python_docx(self, file_path: str, filename: str) -> str:
        """使用python-docx提取.docx文件内容"""
        try:
            doc = DocxDocument(file_path)
            markdown_content = []
            
            # 添加文档标题
            document_title = Path(filename).stem
            markdown_content.append(f"# {document_title}\n")
            
            # 处理段落
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    # 检查是否为标题（基于样式或格式）
                    if paragraph.style.name.startswith('Heading'):
                        # 根据标题级别添加#
                        level = self._get_heading_level(paragraph.style.name)
                        markdown_content.append(f"{'#' * level} {text}\n")
                    else:
                        # 检查是否为列表项
                        if paragraph.style.name.startswith('List'):
                            markdown_content.append(f"- {text}\n")
                        else:
                            # 普通段落
                            markdown_content.append(f"{text}\n")
            
            # 处理表格
            for table in doc.tables:
                table_md = self._convert_table_to_markdown(table)
                if table_md:
                    markdown_content.append(table_md)
            
            result = '\n'.join(markdown_content)
            current_app.logger.info(f"使用python-docx成功转换docx文件: {filename}")
            return result
            
        except Exception as e:
            current_app.logger.error(f"python-docx提取.docx内容失败: {e}")
            return self._generate_doc_conversion_message("DOCX文档解析失败", str(e))
    
    def _extract_doc_content(self, file_path: str, filename: str) -> str:
        """提取.doc文件内容"""
        try:
            # 方法1：使用docx2txt
            try:
                content = docx2txt.process(file_path)
                if content and content.strip():
                    document_title = Path(filename).stem
                    markdown_content = f"# {document_title}\n\n{self._convert_to_markdown(content)}"
                    current_app.logger.info(f"使用docx2txt成功转换doc文件: {filename}")
                    return markdown_content
                else:
                    current_app.logger.warning(f"docx2txt未能提取到内容")
                    return self._extract_doc_content_fallback(file_path, filename)
            except Exception as e:
                error_msg = str(e)
                current_app.logger.warning(f"docx2txt处理失败: {error_msg}")
                
                # 检查是否是文件格式问题
                if "not a zip file" in error_msg.lower() or "bad zip file" in error_msg.lower():
                    # 这是真正的.doc文件，不是.docx文件
                    return self._generate_doc_conversion_message("旧版本的.doc文件格式（非OpenXML）")
                elif "no such file" in error_msg.lower() or "not found" in error_msg.lower():
                    return self._generate_doc_conversion_message("文件不存在或无法访问")
                else:
                    # 尝试fallback方法
                    return self._extract_doc_content_fallback(file_path, filename)
                    
        except Exception as e:
            current_app.logger.error(f"提取.doc内容失败: {e}")
            return self._generate_doc_conversion_message("DOC文档处理失败", str(e))
    
    def _extract_doc_content_fallback(self, file_path: str, filename: str) -> str:
        """使用备用方法提取.doc文件内容"""
        try:
            # 尝试使用python-docx打开.doc文件（有时也能工作）
            doc = DocxDocument(file_path)
            content = []
            
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    content.append(text)
            
            if content:
                document_title = Path(filename).stem
                content_text = '\n'.join(content)
                markdown_content = f"# {document_title}\n\n{self._convert_to_markdown(content_text)}"
                current_app.logger.info(f"使用python-docx fallback成功转换doc文件: {filename}")
                return markdown_content
            else:
                return self._generate_doc_conversion_message("文档内容为空")
            
        except Exception as e:
            error_msg = str(e)
            current_app.logger.error(f"备用方法提取.doc内容失败: {error_msg}")
            
            # 分析错误类型，提供更有用的错误信息
            if "not a zip file" in error_msg.lower():
                return self._generate_doc_conversion_message("旧版本的.doc文件格式")
            elif "not a valid file" in error_msg.lower():
                return self._generate_doc_conversion_message("文件格式不正确或已损坏")
            elif "unsupported format" in error_msg.lower():
                return self._generate_doc_conversion_message("不支持的文件格式")
            else:
                return self._generate_doc_conversion_message("未知错误", error_msg)
    
    def _get_heading_level(self, style_name: str) -> int:
        """根据样式名获取标题级别"""
        if 'Heading 1' in style_name:
            return 2  # 文档标题已经是#，所以从##开始
        elif 'Heading 2' in style_name:
            return 3
        elif 'Heading 3' in style_name:
            return 4
        elif 'Heading 4' in style_name:
            return 5
        elif 'Heading 5' in style_name:
            return 6
        elif 'Heading 6' in style_name:
            return 6
        else:
            return 3  # 默认为三级标题
    
    def _convert_table_to_markdown(self, table) -> str:
        """将Word表格转换为Markdown表格"""
        try:
            if not table.rows:
                return ""
            
            markdown_table = []
            
            # 处理表头
            header_row = table.rows[0]
            header_cells = [cell.text.strip() for cell in header_row.cells]
            if any(header_cells):  # 只有当表头有内容时才添加
                markdown_table.append('| ' + ' | '.join(header_cells) + ' |')
                markdown_table.append('| ' + ' | '.join(['---'] * len(header_cells)) + ' |')
                
                # 处理数据行
                for row in table.rows[1:]:
                    row_cells = [cell.text.strip() for cell in row.cells]
                    markdown_table.append('| ' + ' | '.join(row_cells) + ' |')
            else:
                # 没有明确的表头，所有行都作为数据行
                for row in table.rows:
                    row_cells = [cell.text.strip() for cell in row.cells]
                    markdown_table.append('| ' + ' | '.join(row_cells) + ' |')
            
            return '\n'.join(markdown_table) + '\n\n'
            
        except Exception as e:
            current_app.logger.error(f"转换表格失败: {e}")
            return "\n[表格转换失败]\n\n"
    
    def _convert_to_markdown(self, text: str) -> str:
        """将纯文本转换为基本的Markdown格式"""
        if not text:
            return ""
        
        lines = text.split('\n')
        markdown_lines = []
        
        for line in lines:
            line = line.strip()
            if not line:
                markdown_lines.append("")
                continue
            
            # 简单的格式化规则
            if line.isupper() and len(line) < 50:
                # 全大写且较短的行可能是标题
                markdown_lines.append(f"## {line}")
            elif line.endswith(':') and len(line) < 50:
                # 以冒号结尾的行可能是小标题
                markdown_lines.append(f"### {line}")
            elif line.startswith('•') or line.startswith('-') or line.startswith('*'):
                # 列表项
                markdown_lines.append(f"- {line[1:].strip()}")
            else:
                # 普通段落
                markdown_lines.append(line)
        
        return '\n'.join(markdown_lines)
    
    def _generate_doc_conversion_message(self, reason: str, error_detail: str = None) -> str:
        """生成doc文件转换错误的信息"""
        message = f"""# 文档内容提取失败

## 错误原因
{reason}

## 解决方案
1. **推荐方案**: 使用Microsoft Word打开此文件，另存为.docx格式后重新上传
2. **替代方案**: 使用在线文档转换工具将.doc转换为.docx格式
3. **备用方案**: 将文档内容复制到新的Word文档中保存为.docx格式

## 技术说明
- 当前系统支持现代的.docx格式（Office 2007及以上版本）
- 使用mammoth库进行高质量的markdown转换
- 旧版本的.doc格式（Office 97-2003）需要转换后才能正确解析
- 如果文件实际上是.docx格式但扩展名为.doc，请直接重命名为.docx

## 支持的文档格式
- ✅ .docx (推荐，支持完整格式转换)
- ✅ .md (Markdown)
- ✅ .pdf (通过OCR)
- ❌ .doc (需要转换)

---
*如果问题持续存在，请检查文档是否完整且未损坏。*"""
        
        if error_detail:
            message += f"\n\n## 详细错误信息\n```\n{error_detail}\n```"
        
        return message
    
    def _check_and_create_knowledge_base(self, document: Document):
        """检查并创建知识库（如果是首次上传）"""
        try:
            # 导入知识库服务
            from services.knowledge_base_service import knowledge_base_service
            
            project_id = document.project_id
            user_id = document.upload_by
            
            # 检查是否是首次上传（即当前文档是项目的第一个完成的文档）
            completed_docs_count = Document.query.filter_by(
                project_id=project_id,
                status=DocumentStatus.COMPLETED
            ).count()
            
            current_app.logger.info(f"项目 {project_id} 已完成文档数量: {completed_docs_count}")
            
            # 如果是首次上传完成的文档，创建知识库
            if completed_docs_count == 1:  # 当前文档是第一个完成的文档
                current_app.logger.info(f"项目 {project_id} 首次上传文档完成，开始创建知识库")
                
                # 创建知识库
                dataset_id = knowledge_base_service.create_knowledge_base_for_project(
                    project_id=project_id,
                    user_id=user_id
                )
                
                if dataset_id:
                    current_app.logger.info(f"成功为项目 {project_id} 创建知识库，dataset_id: {dataset_id}")
                    
                    # 上传当前文档到知识库
                    self._upload_document_to_knowledge_base(document)
                else:
                    current_app.logger.error(f"为项目 {project_id} 创建知识库失败")
            else:
                current_app.logger.info(f"项目 {project_id} 已有其他完成的文档，知识库应该已存在")
                
                # 上传当前文档到已存在的知识库
                self._upload_document_to_knowledge_base(document)
                
        except Exception as e:
            current_app.logger.error(f"检查和创建知识库失败: {e}")
            # 不影响文档处理的主流程，只记录错误
    
    def _upload_document_to_knowledge_base(self, document: Document):
        """将文档上传到知识库"""
        try:
            from services.knowledge_base_service import knowledge_base_service
            
            # 检查项目是否有知识库
            project = document.project
            if not project or not project.dataset_id:
                current_app.logger.warning(f"项目 {document.project_id} 没有知识库，跳过文档上传")
                return
            
            current_app.logger.info(f"开始上传文档 {document.name} 到知识库 {project.dataset_id}")
            
            # 上传文档到知识库
            success = knowledge_base_service.upload_document_to_knowledge_base(
                project_id=document.project_id,
                document_id=document.id
            )
            
            if success:
                current_app.logger.info(f"成功上传文档 {document.name} 到知识库")
            else:
                current_app.logger.error(f"上传文档 {document.name} 到知识库失败")
                
        except Exception as e:
            current_app.logger.error(f"上传文档到知识库异常: {e}")
            # 不影响文档处理的主流程，只记录错误

# 全局文档处理器实例
document_processor = DocumentProcessor()
